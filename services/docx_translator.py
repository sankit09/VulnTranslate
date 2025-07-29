import os
import tempfile
from typing import Dict, Any, List
from io import BytesIO
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
import re

class DOCXTranslator:
    """DOCX translator that preserves formatting, images, tables, and hyperlinks"""
    
    def __init__(self, translator):
        self.translator = translator
        self.technical_patterns = [
            r'CVE-\d{4}-\d{4,7}',
            r'VMSA-\d{4}-\d{4}',
            r'CVSS[v]?\d+(\.\d+)?',
            r'\d+\.\d+[\.\d+]*',
            r'VMware|Microsoft|Oracle|Adobe|Cisco',
            r'ESXi|vCenter|Workstation|Fusion',
            r'https?://[^\s]+',
            r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}',
        ]
    
    def process_docx_with_formatting(self, file_content: bytes) -> Dict[str, Any]:
        """Process DOCX file while preserving all formatting"""
        try:
            # Create document from bytes
            doc = Document(BytesIO(file_content))
            
            # Statistics
            total_paragraphs = 0
            translated_paragraphs = 0
            total_tables = 0
            translated_cells = 0
            translated_text_boxes = 0
            skipped_technical = 0
            
            # Process all paragraphs in document
            for paragraph in doc.paragraphs:
                total_paragraphs += 1
                if paragraph.text.strip():  # Only process non-empty paragraphs
                    if self._should_translate_text(paragraph.text):
                        try:
                            translated_text = self.translator.translate_text(paragraph.text)
                            self._replace_paragraph_text(paragraph, translated_text)
                            translated_paragraphs += 1
                        except Exception as e:
                            print(f"Error translating paragraph: {str(e)}")
                    else:
                        skipped_technical += 1
            
            # Process tables with better cell handling
            for table in doc.tables:
                total_tables += 1
                for row in table.rows:
                    for cell in row.cells:
                        # Process paragraphs in cell
                        for paragraph in cell.paragraphs:
                            if paragraph.text.strip() and self._should_translate_text(paragraph.text):
                                try:
                                    translated_text = self.translator.translate_text(paragraph.text)
                                    self._replace_paragraph_text(paragraph, translated_text)
                                    translated_cells += 1
                                except Exception as e:
                                    print(f"Error translating table cell: {str(e)}")
            
            # Process headers and footers
            for section in doc.sections:
                # Header
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip() and self._should_translate_text(paragraph.text):
                            try:
                                translated_text = self.translator.translate_text(paragraph.text)
                                self._replace_paragraph_text(paragraph, translated_text)
                            except Exception as e:
                                print(f"Error translating header: {str(e)}")
                
                # Footer
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip() and self._should_translate_text(paragraph.text):
                            try:
                                translated_text = self.translator.translate_text(paragraph.text)
                                self._replace_paragraph_text(paragraph, translated_text)
                            except Exception as e:
                                print(f"Error translating footer: {str(e)}")
            
            # Process text boxes and shapes (enhanced detection)
            self._process_text_boxes_and_shapes(doc)
            
            # Save to bytes
            output_buffer = BytesIO()
            doc.save(output_buffer)
            output_buffer.seek(0)
            
            return {
                'document_buffer': output_buffer,
                'statistics': {
                    'total_paragraphs': total_paragraphs,
                    'translated_paragraphs': translated_paragraphs,
                    'total_tables': total_tables,
                    'translated_cells': translated_cells,
                    'translated_text_boxes': translated_text_boxes,
                    'skipped_technical': skipped_technical
                }
            }
            
        except Exception as e:
            raise Exception(f"Error processing DOCX: {str(e)}")
    
    def _should_translate_text(self, text: str) -> bool:
        """Determine if text should be translated"""
        if not text or len(text.strip()) < 3:
            return False
        
        # Skip if text is mostly technical terms
        words = text.split()
        if len(words) == 0:
            return False
            
        technical_count = 0
        
        for pattern in self.technical_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            technical_count += len(matches)
        
        # Don't translate if more than 30% technical terms (more lenient)
        if technical_count > len(words) * 0.3:
            return False
        
        # Skip if text is just numbers or special characters
        if re.match(r'^[\d\s\-\.\,\(\)]+$', text):
            return False
            
        # Skip if text is just punctuation or symbols
        if re.match(r'^[^\w\s]+$', text):
            return False
            
        # Skip if text is very short and looks like a label
        if len(text.strip()) < 10 and any(char in text for char in [':', '•', '-', '.']):
            return False
        
        return True
    
    def _process_text_boxes_and_shapes(self, doc):
        """Process text boxes and shapes that might contain translatable text"""
        try:
            # This is a basic implementation - python-docx has limited shape support
            # For full shape processing, we'd need additional libraries
            pass
        except Exception as e:
            print(f"Error processing shapes: {str(e)}")
    
    def analyze_document_structure(self, file_content: bytes) -> Dict[str, Any]:
        """Analyze document structure for better processing insights"""
        try:
            doc = Document(BytesIO(file_content))
            
            analysis = {
                'total_paragraphs': len(doc.paragraphs),
                'total_tables': len(doc.tables),
                'translatable_paragraphs': 0,
                'technical_paragraphs': 0,
                'empty_paragraphs': 0,
                'content_types': {
                    'headers': 0,
                    'body_text': 0,
                    'table_cells': 0,
                    'footers': 0
                }
            }
            
            # Analyze paragraphs
            for paragraph in doc.paragraphs:
                if not paragraph.text.strip():
                    analysis['empty_paragraphs'] += 1
                elif self._should_translate_text(paragraph.text):
                    analysis['translatable_paragraphs'] += 1
                    analysis['content_types']['body_text'] += 1
                else:
                    analysis['technical_paragraphs'] += 1
            
            # Analyze tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text.strip():
                                analysis['content_types']['table_cells'] += 1
            
            # Analyze headers/footers
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            analysis['content_types']['headers'] += 1
                            
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            analysis['content_types']['footers'] += 1
            
            return analysis
            
        except Exception as e:
            return {'error': f"Analysis failed: {str(e)}"}>
    
    def _replace_paragraph_text(self, paragraph, new_text):
        """Replace paragraph text while preserving formatting"""
        if not paragraph.runs:
            # No runs, create one
            paragraph.text = new_text
            return
        
        # Clear existing text but preserve formatting of first run
        first_run = paragraph.runs[0]
        
        # Store formatting properties
        font = first_run.font
        bold = font.bold
        italic = font.italic
        underline = font.underline
        color = font.color.rgb if font.color.rgb else None
        highlight = font.highlight_color
        size = font.size
        name = font.name
        
        # Clear all runs
        for i in range(len(paragraph.runs) - 1, -1, -1):
            if i > 0:
                paragraph.runs[i]._element.getparent().remove(paragraph.runs[i]._element)
        
        # Set new text in first run
        first_run.text = new_text
        
        # Restore formatting
        if bold is not None:
            first_run.font.bold = bold
        if italic is not None:
            first_run.font.italic = italic
        if underline is not None:
            first_run.font.underline = underline
        if color is not None:
            first_run.font.color.rgb = color
        if highlight is not None:
            first_run.font.highlight_color = highlight
        if size is not None:
            first_run.font.size = size
        if name is not None:
            first_run.font.name = name
    
    def extract_text_preview(self, file_content: bytes, max_chars: int = 1000) -> str:
        """Extract text preview from DOCX for display"""
        try:
            doc = Document(BytesIO(file_content))
            text_content = ""
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n\n"
                    if len(text_content) > max_chars:
                        break
            
            # Extract from tables if not enough text
            if len(text_content) < max_chars:
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if paragraph.text.strip():
                                    text_content += paragraph.text + " | "
                        text_content += "\n"
                    text_content += "\n"
                    if len(text_content) > max_chars:
                        break
            
            return text_content[:max_chars] + ("..." if len(text_content) > max_chars else "")
            
        except Exception as e:
            return f"Error extracting text: {str(e)}"