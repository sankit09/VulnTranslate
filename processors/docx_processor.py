"""
DOCX Document Processor
Implements IDocumentProcessor interface for DOCX files using python-docx
"""

import tempfile
import io
from typing import Dict, Any, List
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_COLOR_INDEX

from core.interfaces import IDocumentProcessor
from core.models import DocumentContent, DocumentType
from core.exceptions import ProcessingError, UnsupportedFormatError


class DOCXProcessor(IDocumentProcessor):
    """DOCX document processor with full format preservation"""
    
    def __init__(self):
        self.supported_extensions = ['.docx']
        self.preserve_formatting = True
        self.preserve_images = True
        self.preserve_tables = True
        self.preserve_hyperlinks = True

    def can_process(self, file_extension: str) -> bool:
        """Check if processor can handle DOCX files"""
        return file_extension.lower() in self.supported_extensions

    def extract_content(self, file_content: bytes) -> Dict[str, Any]:
        """Extract translatable content from DOCX while preserving structure"""
        try:
            # Load document from bytes
            doc = Document(io.BytesIO(file_content))
            
            # CRITICAL FIX: Remove first page content BEFORE processing for translation
            self._remove_first_page_content_early(doc)
            
            content_blocks = []
            tables = []
            hyperlinks = []
            images = []
            
            # Process ALL document paragraphs (including empty ones for structure)
            for para_idx, paragraph in enumerate(doc.paragraphs):
                content_block = self._process_paragraph(paragraph, para_idx)
                # Include ALL paragraphs, even empty ones, to maintain structure
                if content_block is not None:
                    content_blocks.append(content_block)
            
            # Process text boxes and shapes
            for shape_idx, shape in enumerate(self._get_document_shapes(doc)):
                shape_content = self._process_shape_text(shape, shape_idx)
                if shape_content:
                    content_blocks.extend(shape_content)
            
            # Process tables
            for table_idx, table in enumerate(doc.tables):
                table_data = self._process_table(table, table_idx)
                tables.append(table_data)
                
                # Add table cell content to content blocks
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for para_idx, paragraph in enumerate(cell.paragraphs):
                            content_block = self._process_paragraph(
                                paragraph, 
                                f"table_{table_idx}_row_{row_idx}_cell_{cell_idx}_para_{para_idx}"
                            )
                            if content_block:
                                content_block['location'] = 'table'
                                content_block['table_info'] = {
                                    'table_idx': table_idx,
                                    'row_idx': row_idx,
                                    'cell_idx': cell_idx
                                }
                                content_blocks.append(content_block)
            
            # Extract hyperlinks
            for paragraph in doc.paragraphs:
                hyperlinks.extend(self._extract_hyperlinks(paragraph))
            
            # Count statistics
            total_paragraphs = len(doc.paragraphs)
            translatable_paragraphs = len([b for b in content_blocks if b.get('translatable', False)])
            technical_paragraphs = total_paragraphs - translatable_paragraphs
            
            document_content = DocumentContent(
                content_blocks=content_blocks,
                metadata={
                    'original_file_size': len(file_content),
                    'document_sections': len(doc.sections),
                    'total_elements': len(content_blocks)
                },
                document_type=DocumentType.DOCX,
                total_paragraphs=total_paragraphs,
                translatable_paragraphs=translatable_paragraphs,
                technical_paragraphs=technical_paragraphs,
                tables=tables,
                hyperlinks=hyperlinks,
                images=images
            )
            
            return {
                'document_content': document_content,
                'original_document': doc,
                'success': True
            }
            
        except Exception as e:
            raise ProcessingError(
                f"Failed to extract content from DOCX: {str(e)}",
                error_code="DOCX_EXTRACTION_FAILED"
            )

    def reconstruct_document(self, content: Dict[str, Any], translations: Dict[str, str]) -> bytes:
        """Reconstruct DOCX document with Japanese first page image and translated content"""
        try:
            original_doc = content.get('original_document')
            if not original_doc:
                raise ProcessingError("Original document not found in content")
            
            # Create a copy of the document for modification
            doc_copy = self._deep_copy_document(original_doc)
            
            # Insert Japanese template image at the beginning
            self._insert_japanese_first_page_image(doc_copy)
            
            # Apply translations to paragraphs (skipping static content)
            self._apply_translations_to_paragraphs(doc_copy, translations)
            
            # Apply translations to tables
            self._apply_translations_to_tables(doc_copy, translations)
            
            # Save to bytes
            output_buffer = io.BytesIO()
            doc_copy.save(output_buffer)
            output_buffer.seek(0)
            
            return output_buffer.getvalue()
            
        except Exception as e:
            raise ProcessingError(
                f"Failed to reconstruct DOCX document: {str(e)}",
                error_code="DOCX_RECONSTRUCTION_FAILED"
            )

    def _process_paragraph(self, paragraph, para_id) -> Dict[str, Any]:
        """Process individual paragraph and extract content - INCLUDE ALL paragraphs"""
        text = paragraph.text.strip()
        
        # CRITICAL: Include ALL paragraphs for complete document coverage
        # Empty paragraphs maintain document structure and spacing
        
        # Check if this is static first-page content that should not be translated
        static_translations = self._get_static_translations()
        if text in static_translations:
            translatable = False  # Mark static content as non-translatable
        else:
            # Determine if paragraph is translatable (even empty ones need processing)
            translatable = self._is_translatable_text(text) if text else False
        
        content_block = {
            'id': str(para_id),
            'text': text if text else '',  # Include empty paragraphs
            'translatable': translatable,
            'location': 'body',
            'formatting': self._extract_paragraph_formatting(paragraph),
            'runs': self._extract_run_formatting(paragraph),
            'is_empty': not bool(text),  # Track empty paragraphs for reconstruction
            'structure_important': True  # All paragraphs maintain document structure
        }
        
        return content_block

    def _get_document_shapes(self, doc):
        """Extract shapes and text boxes from document"""
        shapes = []
        try:
            # Get shapes from document sections
            for section in doc.sections:
                # Access the underlying XML to find text boxes and shapes
                for element in section._sectPr.xpath('.//w:drawing'):
                    shapes.append(element)
        except Exception:
            pass  # If we can't access shapes, continue without them
        return shapes

    def _process_shape_text(self, shape, shape_idx):
        """Extract text from shapes and text boxes"""
        try:
            content_blocks = []
            # This is a simplified version - full implementation would need
            # more complex XML parsing to extract text from shapes
            # For now, we'll focus on improving paragraph and table processing
            return content_blocks
        except Exception:
            return []

    def _process_table(self, table, table_idx) -> Dict[str, Any]:
        """Process table structure and content"""
        table_data = {
            'index': table_idx,
            'rows': len(table.rows),
            'columns': len(table.columns) if table.rows else 0,
            'cells': []
        }
        
        for row_idx, row in enumerate(table.rows):
            row_data = []
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                cell_data = {
                    'text': cell_text,
                    'translatable': self._is_translatable_text(cell_text),
                    'row': row_idx,
                    'column': cell_idx
                }
                row_data.append(cell_data)
            table_data['cells'].append(row_data)
        
        return table_data

    def _extract_hyperlinks(self, paragraph) -> List[Dict[str, Any]]:
        """Extract hyperlinks from paragraph"""
        hyperlinks = []
        
        # Note: python-docx doesn't have direct hyperlink access
        # This is a simplified implementation
        if hasattr(paragraph, '_element'):
            for run in paragraph.runs:
                if hasattr(run, '_element'):
                    # Check for hyperlink relationships
                    # This would need more complex implementation for full hyperlink extraction
                    pass
        
        return hyperlinks

    def _extract_paragraph_formatting(self, paragraph) -> Dict[str, Any]:
        """Extract paragraph-level formatting"""
        formatting = {}
        
        if paragraph.style:
            formatting['style'] = paragraph.style.name
        
        if hasattr(paragraph, 'alignment') and paragraph.alignment:
            formatting['alignment'] = str(paragraph.alignment)
        
        return formatting

    def _extract_run_formatting(self, paragraph) -> List[Dict[str, Any]]:
        """Extract run-level formatting for each text run"""
        runs = []
        
        for run in paragraph.runs:
            run_data = {
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_size': run.font.size.pt if run.font.size else None,
                'font_name': run.font.name,
                'font_color': self._get_font_color(run)
            }
            runs.append(run_data)
        
        return runs

    def _get_font_color(self, run) -> str:
        """Extract font color from run"""
        try:
            if run.font.color and run.font.color.rgb:
                return str(run.font.color.rgb)
        except:
            pass
        return None



    def _apply_static_translation(self, text: str) -> str:
        """Apply static translations to known recurring content"""
        static_translations = self._get_static_translations()
        
        # Check for exact matches first
        if text.strip() in static_translations:
            return static_translations[text.strip()]
        
        # Check for partial matches in longer sentences
        translated_text = text
        for english, japanese in static_translations.items():
            if english.lower() in text.lower():
                translated_text = translated_text.replace(english, japanese)
        
        return translated_text

    def _insert_japanese_first_page_image(self, doc):
        """Insert the pre-translated Japanese first page image and remove original first page content"""
        try:
            from docx.shared import Inches
            from docx.enum.text import WD_BREAK
            import os
            
            # Path to the Japanese template image
            image_path = "japanese_first_page_template.png"
            
            if not os.path.exists(image_path):
                print(f"Warning: Japanese template image not found at {image_path}")
                return
            
            # First, identify and remove all first page content
            self._remove_first_page_content(doc)
            
            # Insert image at the absolute beginning by manipulating XML directly
            body = doc._element.body
            
            # Create a new paragraph element for the image
            from docx.oxml import parse_xml
            new_para_xml = parse_xml('<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"></w:p>')
            
            # Insert at the very beginning of the body
            body.insert(0, new_para_xml)
            
            # Create paragraph object from the XML element
            from docx.text.paragraph import Paragraph
            image_paragraph = Paragraph(new_para_xml, doc)
            
            # Add the Japanese image to this paragraph
            run = image_paragraph.add_run()
            run.add_picture(image_path, width=Inches(7))
            
            # Add a page break to separate from CVE content
            run.add_break(WD_BREAK.PAGE)
            
            print("Successfully inserted Japanese first page image at document beginning")
            
        except Exception as e:
            print(f"Warning: Could not insert Japanese template image: {e}")
            # Fallback: try simple insertion
            try:
                if doc.paragraphs:
                    first_para = doc.paragraphs[0]
                    first_para.clear()
                    run = first_para.add_run()
                    run.add_picture(image_path, width=Inches(6))
                    run.add_break()
            except Exception:
                pass

    def _remove_first_page_content(self, doc):
        """Remove all original first page content before inserting Japanese image"""
        try:
            # First, let's debug what we have
            print("=== DEBUGGING DOCUMENT CONTENT ===")
            for i, paragraph in enumerate(doc.paragraphs[:30]):  # Check first 30 paragraphs
                text = paragraph.text.strip()
                print(f"Para {i}: '{text[:100]}...' (len: {len(text)})")
            
            # Very aggressive approach - remove everything until we find the first VMware/CVE line
            paragraphs_to_remove = []
            cve_started = False
            
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip().lower()
                
                # Look for the very first line that mentions VMware products or CVE content
                if not cve_started:
                    # This should be the first actual CVE content line
                    if (("vmware" in text and any(product in text for product in ["esxi", "vcenter", "workstation", "fusion"])) or
                        ("vmsa-" in text) or
                        ("cve-" in text and len(text) > 15) or
                        ("脆弱性の詳細" in text) or  # In case it's already translated
                        ("vulnerability details" in text.lower())):
                        
                        print(f"=== CVE CONTENT STARTS AT PARAGRAPH {i} ===")
                        print(f"First CVE line: '{paragraph.text[:100]}...'")
                        cve_started = True
                        break
                    else:
                        # Remove everything before CVE content
                        paragraphs_to_remove.append(paragraph)
                        print(f"Marking paragraph {i} for removal: '{text[:50]}...'")
            
            # Remove all paragraphs before CVE content
            print(f"=== REMOVING {len(paragraphs_to_remove)} PARAGRAPHS ===")
            for paragraph in paragraphs_to_remove:
                try:
                    p = paragraph._element
                    p.getparent().remove(p)
                    print(f"Removed: '{paragraph.text[:50]}...'")
                except Exception as e:
                    print(f"Could not remove paragraph: {e}")
                    
            print("=== FIRST PAGE CONTENT REMOVAL COMPLETE ===")
                    
        except Exception as e:
            print(f"Warning: Could not remove first page content: {e}")

    def _remove_first_page_content_early(self, doc):
        """Remove first page content BEFORE translation to prevent it from being processed"""
        try:
            print("=== EARLY FIRST PAGE REMOVAL (BEFORE TRANSLATION) ===")
            
            # Find paragraphs containing first page marketing content
            first_page_indicators = [
                "as the attack surface expands",
                "attack surface",
                "sophisticated threat actors", 
                "vulnerability management has shifted",
                "proactive and predictive approach",
                "transformation calls for",
                "risk-based methodologies",
                "advanced threat intelligence",
                "broader security architecture",
                "contemporary vulnerability management",
                "organization's distinct threat environment",
                "customize remediation strategies",
                "advanced vulnerability management",
                "avm services",
                "risk-based approach", 
                "asset value",
                "severity of vulnerabilities",
                "threat actors",
                "this is where our proactive",
                "build a robust vulnerability management system",
                "based on a risk-based approach",
                "come to play and help you"
            ]
            
            paragraphs_to_remove = []
            cve_content_started = False
            
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip().lower()
                
                # Look for the start of actual CVE content
                is_cve_content = (
                    ("vmware" in text and any(product in text for product in ["esxi", "vcenter", "workstation", "fusion"])) or
                    ("vmsa-" in text and len(text) > 10) or
                    ("cve-" in text and len(text) > 15) or
                    ("vulnerability detail" in text and len(text) > 10) or
                    ("脆弱性の詳細" in text)  # Japanese translation
                )
                
                if is_cve_content:
                    print(f"CVE content detected at paragraph {i}, stopping early removal")
                    cve_content_started = True
                    break
                
                # Remove first page content before CVE content starts
                if not cve_content_started:
                    # Check if paragraph contains first page indicators
                    contains_first_page_content = any(indicator in text for indicator in first_page_indicators)
                    
                    # Remove if it contains first page content OR if it's before CVE content
                    if contains_first_page_content or len(text) < 50:  # Remove short paragraphs too
                        paragraphs_to_remove.append(paragraph)
                        print(f"Marking early paragraph {i} for removal: '{text[:50]}...'")
            
            # Remove the identified paragraphs
            print(f"Removing {len(paragraphs_to_remove)} first page paragraphs BEFORE translation")
            for paragraph in paragraphs_to_remove:
                try:
                    p = paragraph._element
                    p.getparent().remove(p)
                    print(f"Early removed: '{paragraph.text[:50]}...'")
                except Exception as e:
                    print(f"Could not remove paragraph early: {e}")
                    
            print("=== EARLY FIRST PAGE REMOVAL COMPLETE ===")
                    
        except Exception as e:
            print(f"Warning: Could not remove first page content early: {e}")

    def _get_static_translations(self) -> Dict[str, str]:
        """Get mapping of static English text to Japanese translations for first page content"""
        return {
            # First page static content that should be replaced with image
            "Cyber Security Advisory": "サイバーセキュリティアドバイザリ",
            "Advanced Vulnerability Management": "高度脆弱性管理",
            "AVM": "AVM",
            "Attack surface": "アタックサーフェス",
            "threat actors": "脅威アクター",
            "proactive and predictive approach": "プロアクティブで予測的なアプローチ",
            "risk-based methodologies": "リスクベース手法",
            "advanced threat intelligence": "高度な脅威インテリジェンス",
            "broader security architecture": "より広範なセキュリティアーキテクチャ",
            "Contemporary vulnerability management": "現代の脆弱性管理",
            "organization's distinct threat environment": "組織の独特な脅威環境",
            "customize remediation strategies": "修復戦略のカスタマイズ",
            "Proactive Advanced Vulnerability Management": "プロアクティブな高度脆弱性管理",
            "Risk-Based Approach": "リスクベースアプローチ",
            "Asset Value": "資産価値", 
            "Severity of Vulnerabilities": "脆弱性の深刻度",
            "Threat Environment": "脅威環境",
            "robust vulnerability management system": "堅牢な脆弱性管理システム"
        }

    def _is_translatable_text(self, text: str) -> bool:
        """Determine if text should be translated"""
        if not text or len(text.strip()) < 3:
            return False
        
        # Skip technical patterns
        import re
        technical_patterns = [
            r'^CVE-\d{4}-\d{4,7}$',
            r'^https?://',
            r'^\d+\.\d+[\.\d+]*$',
            r'^[A-Z_][A-Z0-9_]*$',  # Constants
            r'^\w+@\w+\.\w+$'  # Emails
        ]
        
        for pattern in technical_patterns:
            if re.match(pattern, text.strip(), re.IGNORECASE):
                return False
        
        return True

    def _deep_copy_document(self, original_doc):
        """Create a deep copy of the document for modification"""
        # Save original to buffer and reload
        buffer = io.BytesIO()
        original_doc.save(buffer)
        buffer.seek(0)
        return Document(buffer)

    def _apply_translations_to_paragraphs(self, doc, translations: Dict[str, str]):
        """Apply translations to document paragraphs with structured reconstruction"""
        for para_idx, paragraph in enumerate(doc.paragraphs):
            para_id = str(para_idx)
            
            # Apply translation if available (including empty strings for empty paragraphs)
            if para_id in translations:
                translated_text = translations[para_id]
                try:
                    # Use enhanced text replacement that preserves formatting
                    self._replace_paragraph_text(paragraph, translated_text)
                except Exception as e:
                    # Fallback to simple text replacement
                    if translated_text:  # Only set non-empty translations
                        paragraph.text = translated_text
            else:
                # Check for static translations for untranslated paragraphs
                original_text = paragraph.text.strip()
                if original_text and len(original_text) > 0:
                    static_translation = self._apply_static_translation(original_text)
                    if static_translation != original_text:
                        try:
                            self._replace_paragraph_text(paragraph, static_translation)
                        except Exception:
                            paragraph.text = static_translation
            
            # Maintain document structure even for untranslated paragraphs

    def _apply_translations_to_tables(self, doc, translations: Dict[str, str]):
        """Apply translations to table cells"""
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        key = f"table_{table_idx}_row_{row_idx}_cell_{cell_idx}_para_{para_idx}"
                        if key in translations:
                            try:
                                self._replace_paragraph_text(paragraph, translations[key])
                            except Exception as e:
                                # Fallback to simple text replacement
                                paragraph.text = translations[key]
                        else:
                            # Apply static translations to untranslated table content
                            original_text = paragraph.text.strip()
                            if original_text and len(original_text) > 0:
                                static_translation = self._apply_static_translation(original_text)
                                if static_translation != original_text:
                                    try:
                                        self._replace_paragraph_text(paragraph, static_translation)
                                    except Exception:
                                        paragraph.text = static_translation

    def _replace_paragraph_text(self, paragraph, new_text: str):
        """Replace paragraph text while preserving formatting"""
        if not paragraph.runs:
            paragraph.text = new_text
            return
        
        # Store original formatting from the first meaningful run
        original_formatting = None
        for run in paragraph.runs:
            if run.text.strip():  # Only use non-empty runs for formatting
                original_formatting = {
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font_size': run.font.size,
                    'font_name': run.font.name,
                    'font_color': run.font.color.rgb if run.font.color.rgb else None
                }
                break
        
        # Simple and reliable approach: clear and rebuild
        paragraph.clear()
        
        # Add new text with preserved formatting
        new_run = paragraph.add_run(new_text)
        
        if original_formatting:
            try:
                # Apply preserved formatting to new run
                if original_formatting['bold'] is not None:
                    new_run.bold = original_formatting['bold']
                if original_formatting['italic'] is not None:
                    new_run.italic = original_formatting['italic']
                if original_formatting['underline'] is not None:
                    new_run.underline = original_formatting['underline']
                if original_formatting['font_size']:
                    new_run.font.size = original_formatting['font_size']
                if original_formatting['font_name']:
                    new_run.font.name = original_formatting['font_name']
                if original_formatting['font_color']:
                    new_run.font.color.rgb = original_formatting['font_color']
            except Exception:
                # If any formatting application fails, continue with plain text
                pass

    def get_document_statistics(self, file_content: bytes) -> Dict[str, Any]:
        """Get detailed statistics about the document"""
        try:
            doc = Document(io.BytesIO(file_content))
            
            stats = {
                'total_paragraphs': len(doc.paragraphs),
                'total_tables': len(doc.tables),
                'total_sections': len(doc.sections),
                'total_pages': self._estimate_page_count(doc),
                'word_count': self._count_words(doc),
                'character_count': self._count_characters(doc)
            }
            
            return stats
            
        except Exception as e:
            return {'error': str(e)}

    def _estimate_page_count(self, doc) -> int:
        """Estimate page count based on content"""
        # Simple estimation based on paragraph count
        paragraph_count = len(doc.paragraphs)
        table_count = len(doc.tables)
        return max(1, (paragraph_count + table_count * 5) // 30)

    def _count_words(self, doc) -> int:
        """Count total words in document"""
        total_words = 0
        for paragraph in doc.paragraphs:
            total_words += len(paragraph.text.split())
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    total_words += len(cell.text.split())
        
        return total_words

    def _count_characters(self, doc) -> int:
        """Count total characters in document"""
        total_chars = 0
        for paragraph in doc.paragraphs:
            total_chars += len(paragraph.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    total_chars += len(cell.text)
        
        return total_chars